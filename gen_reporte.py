import json
from docx import Document
from docx.shared import Inches, Pt
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from docx.shared import RGBColor
import os
import locale

# Establecer idioma local para fechas


def formatear_fecha_larga():
    hoy = datetime.now()
    dia = hoy.day
    mes_en = hoy.strftime("%B").lower()

    meses = {
        "january": "enero", "february": "febrero", "march": "marzo",
        "april": "abril", "may": "mayo", "june": "junio",
        "july": "julio", "august": "agosto", "september": "septiembre",
        "october": "octubre", "november": "noviembre", "december": "diciembre"
    }

    mes_es = meses.get(mes_en, mes_en)
    año = hoy.year
    return f"{dia} de {mes_es} de {año}"
def generar_nombre_archivo():
    hoy = datetime.now()
    dia = hoy.strftime("%d")
    mes = hoy.strftime("%B").lower().replace(" ", "")  # sin espacios
    return f"reporte de prensa {dia}{mes}.docx"

def generar_reporte_word(resumenes, ruta_logo):
    doc = Document()

    # Tabla para encabezado: logo izquierda, fecha derecha
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    row = table.rows[0]
    row.cells[0].width = Inches(3)
    row.cells[1].width = Inches(3)

    # Logo (ajusta tamaño según tu imagen real)
    if os.path.exists(ruta_logo):
        row.cells[0].paragraphs[0].add_run().add_picture(ruta_logo, width=Inches(1.2))

    # Fecha a la derecha
    p_fecha = row.cells[1].paragraphs[0]
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_fecha = p_fecha.add_run(formatear_fecha_larga())
    run_fecha.font.name = "Arial"
    run_fecha.font.size = Pt(17)

    # "Gerencia" centrado debajo
    p_gg = doc.add_paragraph()
    p_gg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_gg = p_gg.add_run("Gerencia de Asuntos Económicos Internacionales")
    run_gg.font.name = "Arial"
    run_gg.font.size = Pt(17)

    # Línea divisoria
    p_linea = doc.add_paragraph()
    p_linea.paragraph_format.space_before = Pt(5)
    run_linea = p_linea.add_run("_____________________________________________________________")
    run_linea.font.color.rgb = RGBColor(0, 0, 0)

    # Título principal
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run("REPORTE DE PRENSA")
    run_titulo.bold = True
    run_titulo.font.name = "Arial"
    run_titulo.font.size = Pt(28.5)

    doc.add_paragraph()

    # Agregar resúmenes
    for item in resumenes:
        # Título de la noticia
        p_titulo = doc.add_paragraph()
        p_titulo.paragraph_format.space_before = Pt(0)
        p_titulo.paragraph_format.space_after = Pt(0)
        run_titulo = p_titulo.add_run(item["titulo"])
        run_titulo.bold = True
        run_titulo.font.name = "Times New Roman"
        run_titulo.font.size = Pt(11)

        # Resumen de la noticia
        p_resumen = doc.add_paragraph(item["resumen"])
        p_resumen.paragraph_format.space_before = Pt(0)
        p_resumen.paragraph_format.space_after = Pt(0)
        p_resumen.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_resumen = p_resumen.runs[0]
        run_resumen.font.name = "Times New Roman"
        run_resumen.font.size = Pt(11)


        doc.add_paragraph()

    nombre_archivo = generar_nombre_archivo()
    doc.save(nombre_archivo)
    print(f"✅ Documento generado: {nombre_archivo}")

# Cargar resúmenes guardados
with open("resumenes_aprobados.json", "r", encoding="utf-8") as f:
    resumenes = json.load(f)

# Ruta del logo
ruta_logo = "logo_bx.png"  # Ajusta si tu archivo se llama distinto

# Generar el Word
generar_reporte_word(resumenes, ruta_logo)

# Eliminar el archivo temporal de resúmenes